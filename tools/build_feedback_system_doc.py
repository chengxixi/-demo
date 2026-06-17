from __future__ import annotations

import html
import math
import subprocess
from pathlib import Path
from typing import Iterable
from xml.sax.saxutils import escape

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DRAWIO_DIR = ROOT / "drawio"
IMAGE_DIR = ROOT / "images"
OUTPUT_DOCX = ROOT / "反馈管理系统需求设计说明书_优化版.docx"


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(9)
    set_run_font(run)


def set_run_font(run) -> None:
    run.font.name = "PingFang SC"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")


def set_doc_font(doc: Document) -> None:
    styles = doc.styles
    for style_name in ["Normal", "Heading 1", "Heading 2", "Heading 3"]:
        style = styles[style_name]
        style.font.name = "PingFang SC"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    styles["Normal"].font.size = Pt(10.5)
    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 2"].font.size = Pt(13)
    styles["Heading 3"].font.size = Pt(11.5)


def shade_cell(cell, color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    tc_pr.append(shd)


def set_cell_border(cell, color: str = "D9E2F3") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:color"), color)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, True)
        shade_cell(table.rows[0].cells[idx], "DCE6F1")
        set_cell_border(table.rows[0].cells[idx])
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value)
            set_cell_border(cells[idx])
    doc.add_paragraph()


def add_para(doc: Document, text: str, bold_prefix: str | None = None) -> None:
    paragraph = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        run = paragraph.add_run(bold_prefix)
        run.bold = True
        set_run_font(run)
        run = paragraph.add_run(text[len(bold_prefix):])
        set_run_font(run)
    else:
        run = paragraph.add_run(text)
        set_run_font(run)


def add_bullets(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        run = paragraph.add_run(item)
        set_run_font(run)


def mx_node(cell_id: str, value: str, x: int, y: int, w: int, h: int, fill: str = "#ffffff",
            stroke: str = "#5B7DB1", rounded: bool = True, font_size: int = 14) -> str:
    rounded_part = "rounded=1;" if rounded else "rounded=0;"
    style = (
        f"{rounded_part}whiteSpace=wrap;html=1;fillColor={fill};strokeColor={stroke};"
        f"fontSize={font_size};fontStyle=1;arcSize=8;"
    )
    return (
        f'<mxCell id="{cell_id}" value="{mx_value(value)}" style="{style}" vertex="1" parent="1">'
        f'<mxGeometry x="{x}" y="{y}" width="{w}" height="{h}" as="geometry"/></mxCell>'
    )


def mx_text(cell_id: str, value: str, x: int, y: int, w: int, h: int, font_size: int = 18,
            color: str = "#1f2937") -> str:
    style = (
        f"text;html=1;strokeColor=none;fillColor=none;align=center;verticalAlign=middle;"
        f"whiteSpace=wrap;fontSize={font_size};fontStyle=1;fontColor={color};"
    )
    return (
        f'<mxCell id="{cell_id}" value="{mx_value(value)}" style="{style}" vertex="1" parent="1">'
        f'<mxGeometry x="{x}" y="{y}" width="{w}" height="{h}" as="geometry"/></mxCell>'
    )


def mx_edge(cell_id: str, source: str, target: str, label: str = "", color: str = "#666666") -> str:
    style = (
        f"edgeStyle=orthogonalEdgeStyle;rounded=0;orthogonalLoop=1;jettySize=auto;html=1;"
        f"endArrow=block;endFill=1;strokeWidth=2;strokeColor={color};fontSize=12;"
    )
    return (
        f'<mxCell id="{cell_id}" value="{mx_value(label)}" style="{style}" edge="1" parent="1" '
        f'source="{source}" target="{target}"><mxGeometry relative="1" as="geometry"/></mxCell>'
    )


def mx_container(cell_id: str, label: str, x: int, y: int, w: int, h: int, fill: str) -> str:
    style = (
        f"rounded=1;whiteSpace=wrap;html=1;fillColor={fill};strokeColor=#93A4B7;"
        "fontSize=16;fontStyle=1;align=left;verticalAlign=top;spacingLeft=12;spacingTop=8;arcSize=6;"
    )
    return (
        f'<mxCell id="{cell_id}" value="{mx_value(label)}" style="{style}" vertex="1" parent="1">'
        f'<mxGeometry x="{x}" y="{y}" width="{w}" height="{h}" as="geometry"/></mxCell>'
    )


def mx_value(value: str) -> str:
    escaped = html.escape(value)
    return escaped.replace("\\n", "&lt;br&gt;").replace("\n", "&lt;br&gt;")


def wrap_drawio(cells: list[str], width: int, height: int) -> str:
    body = "\n".join(cells)
    return f'''<mxfile host="app.diagrams.net" modified="2026-06-04T00:00:00.000Z" agent="Codex" version="24.7.17">
  <diagram id="feedback-system" name="Page-1">
    <mxGraphModel dx="{width}" dy="{height}" grid="1" gridSize="10" guides="1" tooltips="1" connect="1" arrows="1" fold="1" page="1" pageScale="1" pageWidth="{width}" pageHeight="{height}" math="0" shadow="0">
      <root>
        <mxCell id="0"/>
        <mxCell id="1" parent="0"/>
        {body}
      </root>
    </mxGraphModel>
  </diagram>
</mxfile>'''


def build_architecture_diagram() -> str:
    cells = []
    layers = [
        ("layer1", "数据采集层", 40, 55, "#EAF4FF", [
            ("海外电商ERP", "国内电商ERP", "卖家精灵API\\n（一期深度对接）", "站内信/客服录入", "APP反馈", "线下反馈")
        ]),
        ("layer2", "数据处理层", 200, 55, "#EEF7EA", [
            ("数据清洗", "AI翻译/摘要", "AI去重聚类", "三级分类", "异常分级P0-P3", "竞品指标处理")
        ]),
        ("layer3", "业务应用层", 360, 55, "#FFF4E6", [
            ("反馈工单池", "产品异常处理", "竞品监控", "Q&A库", "案例库", "数据看板")
        ]),
        ("layer4", "对外服务层", 520, 55, "#F3ECFF", [
            ("现有需求池", "事件管理", "PMS/售后", "项目管理", "飞书通知/文档", "计划/生产管理")
        ]),
    ]
    node_ids: list[list[str]] = []
    for layer_id, label, y, x0, fill, groups in layers:
        cells.append(mx_container(layer_id, label, 40, y, 1240, 130, fill))
        ids = []
        for i, text in enumerate(groups[0]):
            node_id = f"{layer_id}_n{i}"
            ids.append(node_id)
            cells.append(mx_node(node_id, text, x0 + i * 190, y + 54, 150, 54, "#FFFFFF", "#5B7DB1", True, 13))
        node_ids.append(ids)
    for r in range(3):
        cells.append(mx_edge(f"e_layer_{r}", node_ids[r][2], node_ids[r + 1][2], ""))
        cells.append(mx_edge(f"e_layer_{r}_b", node_ids[r][4], node_ids[r + 1][4], ""))
    return wrap_drawio(cells, 1320, 700)


def build_interface_diagram() -> str:
    cells = [mx_text("title", "外部系统接口关系图", 0, 20, 1300, 40, 24)]
    cells.append(mx_container("input_group", "输入数据源", 45, 95, 350, 590, "#EAF4FF"))
    cells.append(mx_container("output_group", "下游系统 / 接口预留", 900, 95, 350, 590, "#F3ECFF"))
    cells.append(mx_node("core", "反馈管理系统\\n（一期建设系统）", 535, 330, 230, 92, "#EAF4FF", "#3B6EA8", True, 18))
    inputs = [
        ("海外/国内ERP", "拉取/同步\\n退货、订单、售后"),
        ("卖家精灵API", "一期深度对接\\n竞品/关键词/榜单"),
        ("亚马逊/国内平台", "经ERP或接口\\n评论、Q&A、站内信"),
        ("APP/客服渠道", "录入/导入\\n用户原始反馈"),
        ("线下反馈", "表单录入\\n展会/客户/代理商"),
    ]
    outputs = [
        ("现有需求池", "推送候选需求\\n接口预留"),
        ("事件管理系统", "推送产品异常\\n接口预留"),
        ("项目管理系统", "立项后同步\\n接口预留"),
        ("计划管理系统", "计划节点同步\\n接口预留"),
        ("生产管理系统", "生产/质量闭环\\n接口预留"),
        ("飞书", "通知/报表/文档沉淀\\n一期基础通知"),
    ]
    for i, (name, label) in enumerate(inputs):
        node_id = f"in{i}"
        y = 155 + i * 95
        cells.append(mx_node(node_id, f"{name}\\n{label}", 85, y, 270, 64, "#FFFFFF", "#5B7DB1", True, 13))
    for i, (name, label) in enumerate(outputs):
        node_id = f"out{i}"
        y = 140 + i * 82
        cells.append(mx_node(node_id, f"{name}\\n{label}", 935, y, 280, 60, "#FFFFFF", "#6A4C93", True, 13))
    cells.append(mx_edge("input_to_core", "input_group", "core", "采集 / 拉取 / 录入", "#457B9D"))
    cells.append(mx_edge("core_to_output", "core", "output_group", "推送 / 通知 / 接口预留", "#6A4C93"))
    return wrap_drawio(cells, 1300, 760)


def build_data_flow_diagram() -> str:
    cells = [mx_text("title", "反馈数据主流程图", 0, 20, 1320, 40, 24)]
    feedback = [
        ("f1", "反馈采集", 70, 150),
        ("f2", "AI处理\\n翻译/去重/摘要", 270, 150),
        ("f3", "反馈工单池", 500, 150),
        ("f4", "异常识别\\nP0-P3", 720, 150),
        ("f5", "处理闭环\\n或转需求池", 940, 150),
        ("f6", "数据看板\\n复盘沉淀", 1140, 150),
    ]
    for node_id, label, x, y in feedback:
        cells.append(mx_node(node_id, label, x, y, 150, 70, "#EAF4FF", "#3B6EA8", True, 14))
    for i in range(len(feedback) - 1):
        cells.append(mx_edge(f"ef{i}", feedback[i][0], feedback[i + 1][0]))
    comp = [
        ("c1", "卖家精灵API\\n竞品数据", 150, 400),
        ("c2", "竞品监控\\nASIN/榜单/关键词", 390, 400),
        ("c3", "竞品分析\\n趋势/价格/评价", 650, 400),
        ("c4", "市场机会\\n竞品异常提醒", 910, 400),
    ]
    for node_id, label, x, y in comp:
        cells.append(mx_node(node_id, label, x, y, 180, 70, "#FFF4E6", "#D08A2D", True, 14))
    for i in range(len(comp) - 1):
        cells.append(mx_edge(f"ec{i}", comp[i][0], comp[i + 1][0], "", "#D08A2D"))
    cells.append(mx_edge("ec_to_f3", "c4", "f3", "", "#D08A2D"))
    cells.append(mx_text("note1", "用户反馈主线", 70, 100, 220, 34, 18, "#1f2937"))
    cells.append(mx_text("note2", "竞品数据主线", 150, 350, 220, 34, 18, "#1f2937"))
    return wrap_drawio(cells, 1320, 620)


def build_boundary_diagram() -> str:
    cells = [mx_text("title", "一期/二期建设边界图", 0, 20, 1220, 40, 24)]
    cells.append(mx_container("phase1", "一期建设范围：先让反馈闭环跑起来", 70, 100, 510, 530, "#EAF4FF"))
    cells.append(mx_container("phase2", "二期及接口预留：深度闭环与跨系统联动", 640, 100, 510, 530, "#F3ECFF"))
    p1 = [
        "反馈采集：ERP、客服、APP、站内信、线下录入",
        "卖家精灵API深度对接：竞品、关键词、榜单、评价等",
        "AI处理：翻译、摘要、去重、分类建议、异常预警",
        "工单管理：反馈池、分派、处理、关闭、超时提醒",
        "产品异常处理：P0-P3等级、响应、闭环",
        "基础看板：反馈趋势、异常统计、竞品核心指标",
        "ERP/PMS基础对接与飞书通知"
    ]
    p2 = [
        "需求池深度闭环：需求状态、评分、审批、复盘回写",
        "事件管理闭环：异常事件升级、整改验证、案例库联动",
        "项目管理深度联动：立项、任务、里程碑、评审节点",
        "计划管理联动：迭代计划、资源排期、延期预警",
        "生产管理联动：批次、质量、返工、供应链整改",
        "高阶AI分析：聚类洞察、需求优先级建议、自动报告",
        "跨系统数据资产与统一主数据治理"
    ]
    for i, text in enumerate(p1):
        cells.append(mx_node(f"p1_{i}", text, 110, 170 + i * 65, 430, 46, "#FFFFFF", "#3B6EA8", True, 12))
    for i, text in enumerate(p2):
        cells.append(mx_node(f"p2_{i}", text, 680, 170 + i * 65, 430, 46, "#FFFFFF", "#6A4C93", True, 12))
    return wrap_drawio(cells, 1220, 720)


def write_drawio_files() -> list[Path]:
    DRAWIO_DIR.mkdir(exist_ok=True)
    IMAGE_DIR.mkdir(exist_ok=True)
    diagrams = {
        "反馈管理系统_总体架构图": build_architecture_diagram(),
        "反馈管理系统_外部接口图": build_interface_diagram(),
        "反馈管理系统_数据主流程图": build_data_flow_diagram(),
        "反馈管理系统_一期二期边界图": build_boundary_diagram(),
    }
    paths = []
    for name, content in diagrams.items():
        path = DRAWIO_DIR / f"{name}.drawio"
        path.write_text(content, encoding="utf-8")
        paths.append(path)
    return paths


def export_png(drawio_paths: list[Path]) -> list[Path]:
    image_paths = []
    for source in drawio_paths:
        target = IMAGE_DIR / f"{source.stem}.png"
        subprocess.run(
            ["drawio", "--export", "--format", "png", "--scale", "3", "--output", str(target), str(source)],
            check=True,
            cwd=ROOT,
        )
        image_paths.append(target)
    return image_paths


def add_caption(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(89, 89, 89)
    set_run_font(run)


def add_picture(doc: Document, image_name: str, caption: str) -> None:
    image_path = IMAGE_DIR / image_name
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(9.0))
    add_caption(doc, caption)


def build_docx() -> None:
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11.69)
    section.page_height = Inches(8.27)
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)
    set_doc_font(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("反馈管理系统需求设计说明书（优化版）")
    run.bold = True
    run.font.size = Pt(22)
    set_run_font(run)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("一期重点：反馈采集闭环、卖家精灵 API 深度对接、AI 处理、工单与异常管理")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(89, 89, 89)
    set_run_font(run)

    doc.add_heading("一、文档说明", level=1)
    add_para(doc, "本文档基于《产品优化迭代体系框架V2》、本地《用户反馈与竞争分析流程0429.xlsx》以及现有《反馈管理系统需求设计说明书》优化而成，用于明确第一流程节点“反馈与市场信息收集”的 IT 系统建设范围。")
    add_para(doc, "边界说明：本阶段不重新设计现有需求池，也不展开第二流程节点“反馈分析与需求管理”的 IT MVP；需求池、事件管理、项目管理、计划管理、生产管理均作为下游承接系统或接口预留展示。")
    add_bullets(doc, [
        "一期目标：把用户反馈、售后/退货、APP/客服、线下反馈和卖家精灵竞品数据纳入统一反馈管理系统，形成可跟踪、可预警、可统计、可复盘的闭环。",
        "一期原则：MVP 优先，优先支持高频数据源、关键异常和基础看板；复杂审批、跨系统深度联动放到二期。",
        "图示说明：本版新增 Draw.io 源文件和 PNG 图，Word 中嵌入 PNG，项目 drawio 目录保留可编辑源文件。",
    ])

    doc.add_heading("二、业务流程对齐", level=1)
    add_para(doc, "系统设计需直接承接第一流程节点的三个关键流程：用户反馈信息收集、竞品收集与分析、市场端紧急异常处理。系统不是单纯录入工具，而是将流程中的来源、分类、响应、统计、异常升级和沉淀动作固化到系统中。")
    add_table(doc, ["流程来源", "流程关键要求", "系统承接方式", "一期边界"], [
        ["用户反馈信息收集流程", "全渠道收集、初步分类、技术支持核实、紧急问题 24h 同步、Q&A沉淀。", "反馈工单池、渠道表单、AI翻译摘要、三级分类、处理状态、超时提醒、Q&A库。", "先覆盖主渠道和核心字段，支持人工修正 AI 结果。"],
        ["竞品收集与分析流程", "跟踪 BSR、评分、评论、价格、详情页、新品、促销、广告、卖点和差评集中点。", "卖家精灵 API 深度对接、竞品ASIN档案、指标快照、竞品看板、异常提醒。", "卖家精灵已开 API，放入一期正式建设。"],
        ["市场端紧急异常处理流程", "P0/P1/P2/P3 分级、响应时效、责任部门协同、处理结果归档。", "异常识别、等级判定、责任分派、临时处置、闭环归档、案例库。", "先支持产品/市场相关异常工单和提醒，不做复杂事件系统闭环。"],
    ])

    doc.add_heading("三、系统总体方案", level=1)
    add_para(doc, "系统定位：反馈管理系统是产品快速优化迭代体系的前端数据入口和反馈闭环平台，负责将用户反馈、市场反馈、竞品信息和异常信息统一纳入管理，并向现有需求池、事件管理和后续项目/计划/生产系统提供标准化输出。")
    add_table(doc, ["层级", "一期能力", "二期或预留能力"], [
        ["数据采集层", "海外/国内ERP、客服录入、站内信、APP反馈、线下反馈、卖家精灵API。", "更多平台直连、社媒舆情、更多第三方竞品数据源。"],
        ["数据处理层", "清洗、翻译、摘要、去重、三级分类建议、P0-P3异常建议、竞品指标处理。", "相似问题聚类、需求评分建议、自动洞察和策略建议。"],
        ["业务应用层", "反馈工单池、异常处理、竞品监控、Q&A库、案例库、基础看板。", "跨部门绩效、复盘机制、需求闭环深度联动。"],
        ["对外服务层", "ERP/PMS基础对接、飞书通知、现有需求池导入/接口预留。", "事件管理、项目管理、计划管理、生产管理深度联动。"],
    ])

    doc.add_heading("四、系统架构与接口图示", level=1)
    add_picture(doc, "反馈管理系统_总体架构图.png", "图 1：反馈管理系统总体架构图。用于说明本系统从数据采集、AI处理、业务应用到对外服务的四层边界。")
    add_para(doc, "图 1 的一期范围包括反馈采集、卖家精灵 API、AI基础处理、工单管理、异常预警和基础看板；二期重点是需求池、事件管理、项目/计划/生产系统的深度闭环。")
    add_picture(doc, "反馈管理系统_外部接口图.png", "图 2：外部系统接口关系图。用于说明输入系统、输出系统和接口方向。")
    add_para(doc, "图 2 中卖家精灵 API 为一期正式输入接口；需求池、事件管理、项目管理、计划管理、生产管理为下游输出或接口预留，不在本阶段重建。")
    add_picture(doc, "反馈管理系统_数据主流程图.png", "图 3：反馈数据主流程图。用于说明用户反馈主线与竞品数据主线如何进入工单、异常、需求和复盘。")
    add_para(doc, "图 3 将竞品数据作为独立主线接入竞品监控和竞品分析，再输出市场机会、竞品异常提醒或反馈工单。")
    add_picture(doc, "反馈管理系统_一期二期边界图.png", "图 4：一期/二期建设边界图。用于评审时快速确认当前建设范围和后续预留方向。")

    doc.add_heading("五、核心功能需求", level=1)
    add_table(doc, ["模块", "一期功能", "关键说明"], [
        ["反馈采集", "渠道配置、手工录入、批量导入、ERP/PMS基础同步、附件证据上传。", "支持亚马逊评论、退货、站内信、售后沟通、APP反馈、线下反馈。"],
        ["AI智能处理", "翻译、摘要、关键词提取、去重建议、三级分类建议、异常等级建议。", "所有 AI 结论允许人工修正，并保留原文和修正记录。"],
        ["反馈工单池", "待处理、处理中、待确认、已关闭、转异常、转需求池等状态。", "支撑产品经理、运营、售后、技术支持协同处理。"],
        ["产品异常处理", "P0/P1/P2/P3等级、响应时效、责任部门、临时处置、闭环归档。", "一期先做异常工单闭环，事件管理系统作为后续接口。"],
        ["竞品监控", "卖家精灵API、竞品ASIN档案、核心指标快照、价格/评分/评论/榜单趋势。", "卖家精灵深度对接放入一期，不作为二期延后。"],
        ["Q&A库与案例库", "沉淀标准回复、典型问题、异常案例和处理结论。", "支撑售后口径统一和产品复盘。"],
        ["数据看板", "反馈趋势、渠道分布、产品线TOP问题、异常统计、竞品核心指标。", "先做基础经营看板，深度洞察二期增强。"],
    ])

    doc.add_heading("六、卖家精灵 API 对接设计", level=1)
    add_para(doc, "卖家精灵 API 已开通，应作为一期正式建设的数据源，而不是二期尝试项。对接重点不是简单导入报表，而是把竞品与市场指标持续沉淀到系统中，为竞品异常提醒、市场机会识别和后续需求产生提供依据。")
    add_table(doc, ["数据类型", "建议字段", "用途", "一期要求"], [
        ["竞品ASIN档案", "ASIN、品牌、类目、站点、产品线、价格、评分、评论数、上架时间、链接。", "建立竞品基础档案与监控对象。", "支持手工维护和 API 更新。"],
        ["榜单/排名数据", "BSR、类目排名、关键词排名、采集时间、变化幅度。", "识别市场热度、竞品波动和异常变化。", "支持日/周同步和趋势展示。"],
        ["评论与差评数据", "评分、评论原文、时间、主题、差评原因、情绪倾向。", "识别竞品痛点和机会点。", "支持 AI 摘要和分类建议。"],
        ["关键词与搜索数据", "关键词、搜索量、排名、竞品覆盖、趋势。", "辅助市场机会识别和产品卖点优化。", "先纳入核心关键词池。"],
        ["价格与促销数据", "售价、折扣、优惠券、促销活动、变化时间。", "监控价格策略和促销异常。", "支持价格变化提醒。"],
    ])
    add_para(doc, "待确认：卖家精灵 API 的具体 endpoint、鉴权方式、调用频率限制、可返回字段、历史数据回溯范围和费用限制，需要由 IT 与卖家精灵 API 文档进一步确认。")

    doc.add_heading("七、数据对象与字段", level=1)
    add_table(doc, ["数据对象", "核心字段", "说明"], [
        ["反馈记录", "反馈ID、来源渠道、站点、ASIN/SKU、产品线、内部型号、原文、翻译、摘要、附件、反馈时间。", "所有用户与市场反馈的原始载体。"],
        ["反馈工单", "工单ID、反馈ID、状态、责任部门、责任人、优先级、处理动作、完成时限、关闭原因。", "支撑跨部门处理闭环。"],
        ["分类标签", "一级分类、二级分类、三级分类、适用产品线、启用状态、维护人。", "对应流程中的反馈分类规范。"],
        ["异常记录", "异常ID、异常等级、触发来源、影响范围、临时处置、根因、整改、验证、关闭时间。", "承接 P0-P3 异常管理。"],
        ["竞品档案", "竞品ASIN、品牌、类目、站点、产品线、链接、监控状态、负责人。", "竞品监控的主数据。"],
        ["竞品指标快照", "ASIN、采集时间、价格、评分、评论数、BSR、关键词排名、促销状态。", "用于趋势分析和异常提醒。"],
        ["Q&A/案例", "问题、标准回复、适用产品、来源工单、处理结论、版本、生效状态。", "用于售后回复和知识沉淀。"],
    ])

    doc.add_heading("八、页面与交互", level=1)
    add_table(doc, ["页面/模块", "核心交互", "一期验收重点"], [
        ["反馈录入与导入", "选择渠道、填写产品信息、上传附件、导入Excel、自动生成工单。", "字段完整、可追溯原始反馈。"],
        ["反馈工单池", "筛选、分派、处理、转异常、转需求池、关闭、批量操作。", "状态清晰，责任人和时效明确。"],
        ["工单详情页", "展示原文、AI结果、分类、处理记录、附件证据、关联Q&A/案例。", "人工可修正 AI 分类和摘要。"],
        ["异常处理页", "等级判定、响应计时、责任归属、处置记录、验证关闭。", "P0-P3 响应时效可监控。"],
        ["竞品监控页", "竞品档案、指标趋势、价格变化、评分评论变化、关键词变化。", "卖家精灵数据可稳定展示。"],
        ["数据看板", "产品线、渠道、分类、异常、竞品指标、TOP问题。", "支持周/月复盘。"],
        ["系统配置", "渠道、产品线、分类标签、角色权限、提醒规则、接口配置。", "配置可维护，不依赖开发改字段。"],
    ])

    doc.add_heading("九、接口与集成", level=1)
    add_table(doc, ["接口对象", "方向", "一期/二期", "说明"], [
        ["海外/国内ERP", "拉取/同步", "一期", "获取订单、退货、售后、渠道和产品基础数据。"],
        ["PMS/售后相关系统", "拉取/推送", "一期基础", "用于售后反馈、技术支持核实和处理结果同步。"],
        ["卖家精灵API", "拉取", "一期深度", "获取竞品、关键词、榜单、评价、价格促销等数据。"],
        ["飞书", "推送/沉淀", "一期基础", "推送异常提醒、周/月报摘要、沉淀文档或表格。"],
        ["现有需求池", "推送/接口预留", "二期深化", "本阶段只展示下游承接关系，不重建需求池。"],
        ["事件管理系统", "推送/接口预留", "二期深化", "产品异常事件后续可进入事件管理闭环。"],
        ["项目管理系统", "双向同步预留", "二期", "需求立项后同步任务、里程碑和评审节点。"],
        ["计划管理系统", "双向同步预留", "二期", "同步迭代计划、资源排期和延期预警。"],
        ["生产管理系统", "双向同步预留", "二期", "同步批次、质量、返工和供应链整改信息。"],
    ])

    doc.add_heading("十、权限与安全", level=1)
    add_para(doc, "权限设计采用角色 + 产品线/品牌数据范围的组合方式。默认用户只能查看和处理自己职责范围内的数据；跨品牌、跨产品线、跨部门的数据访问需由管理员配置。")
    add_table(doc, ["角色", "权限范围", "关键限制"], [
        ["管理员", "系统配置、角色权限、接口配置、全量数据查看。", "仅限授权人员，所有配置变更记录日志。"],
        ["品牌/产品经理", "负责品牌或产品线的反馈、异常、竞品和看板。", "可转需求池，但不替代现有需求池审批。"],
        ["电商/运营", "负责渠道反馈、竞品监控、平台异常和处理协同。", "仅查看授权站点和店铺数据。"],
        ["售后/客服", "录入和处理售后反馈、维护Q&A、查看处理结论。", "不查看无关竞品和管理层看板。"],
        ["研发/技术支持", "查看产品相关反馈、异常、根因分析和整改记录。", "只处理分派事项。"],
        ["管理层", "查看汇总看板、异常趋势和重点问题。", "默认不参与日常工单编辑。"],
    ])
    add_bullets(doc, [
        "敏感信息保护：系统不得在文档、日志或报表中暴露 API Secret、OAuth token、授权二维码或账号凭据。",
        "数据追溯：所有接口同步、AI处理、人工修正、状态变更和关闭动作必须记录操作者、时间和变更前后内容。",
        "AI安全：AI输出作为建议，不自动关闭工单、不自动定责、不自动生成对外回复，关键结论需人工确认。",
    ])

    doc.add_heading("十一、实施计划与验收标准", level=1)
    add_table(doc, ["阶段", "建设内容", "验收标准"], [
        ["一期MVP", "反馈采集、卖家精灵API、AI处理、工单池、异常预警、基础看板、ERP/PMS基础对接、飞书通知。", "能用真实渠道数据生成工单；卖家精灵数据稳定入库；P0-P3异常可跟踪；周/月报可输出。"],
        ["一期优化", "补充字段、优化分类标签、完善超时提醒、完善竞品看板和处理闭环。", "试运行问题有记录和版本优化；核心用户可完成日常流程。"],
        ["二期深化", "需求池、事件管理、项目管理、计划管理、生产管理深度联动，高阶AI分析。", "跨系统状态可同步，需求/异常/项目/生产数据形成闭环。"],
    ])
    add_para(doc, "建议一期试运行验收用例：")
    add_bullets(doc, [
        "普通用户反馈：亚马逊评论或售后反馈进入系统，经 AI 摘要和三级分类后生成工单，并完成处理关闭。",
        "产品异常反馈：退货或投诉触发 P1/P2 异常建议，责任人处理、验证并归档案例。",
        "竞品监控：卖家精灵 API 同步目标 ASIN 指标，在竞品看板展示价格、评分、评论和榜单变化。",
        "下游承接：经人工确认的高价值反馈可导出或推送至现有需求池，需求池系统继续承接后续管理。",
    ])

    doc.save(OUTPUT_DOCX)


def main() -> None:
    drawio_paths = write_drawio_files()
    export_png(drawio_paths)
    build_docx()
    print(str(OUTPUT_DOCX))


if __name__ == "__main__":
    main()
